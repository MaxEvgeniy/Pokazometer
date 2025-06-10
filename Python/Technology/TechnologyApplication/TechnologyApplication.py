import fitz  # PyMuPDF
import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import messagebox, ttk
from PIL import Image, ImageTk
import io
import pandas as pd


# --- Настройки ---
folder_path = r'C:\PDF\1\Rename'
template_docx_path = os.path.join(folder_path, "template.docx")
excel_file = os.path.join(folder_path, "KeyWords.xlsx")


# --- Регионы A4 ---
rects_A4 = [
    fitz.Rect(245, 676, 576, 709),   # Обозначение детали
    fitz.Rect(245, 720, 434, 750),   # Наименование детали
    fitz.Rect(245, 788, 434, 825),   # Материал детали
    fitz.Rect(105, 744, 170, 757),   # Разработал
    fitz.Rect(212, 744, 240, 757)    # Дата чертежа
]

# --- Смещения для A3 ---
offset_x = 585
offset_y = 1
rects_A3 = [fitz.Rect(r.x0 + offset_x, r.y0 + offset_y, r.x1 + offset_x, r.y1 + offset_y) for r in rects_A4]


def normalize_dashes_and_spaces(text):
    return re.sub(r'[\u00AD\u2010\u2011\u2012\u2013\u2014\u2015]', '-', text).strip()


def extract_text_from_rect(page, rect):
    lines = [line.strip() for line in page.get_text("text", clip=rect).splitlines() if line.strip()]
    return lines[0] if lines else ""


def clean_filename(name):
    name = re.sub(r'\s+', ' ', name).strip()
    return re.sub(r'[\\/*?:"<>|]', '', name)


class App:
    def __init__(self, root, pdf_path, обозначение, наименование, материал, разработал, дата):
        self.root = root
        self.pdf_path = pdf_path
        self.обозначение = обозначение
        self.наименование = наименование
        self.материал = материал
        self.разработал = разработал
        self.дата = дата
        self.выбранный_тип = None
        self.выбранные_операции = []
        self.selected_image_area = None
        self.tk_img_ref = None
        self.zoom_level = 0.7

        # --- GUI ---
        root.title(f"{self.обозначение} {self.наименование}")
        root.geometry("1200x800")

        # Левая часть: превью PDF
        frame_left = tk.Frame(root)
        frame_left.pack(side="left", fill="both", expand=True)

        canvas_frame = tk.Frame(frame_left)
        canvas_frame.pack(fill="both", expand=True)

        self.canvas = tk.Canvas(canvas_frame, bg="white")
        h_scrollbar = tk.Scrollbar(canvas_frame, orient="horizontal", command=self.canvas.xview)
        v_scrollbar = tk.Scrollbar(canvas_frame, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(xscrollcommand=h_scrollbar.set, yscrollcommand=v_scrollbar.set)

        h_scrollbar.pack(side="bottom", fill="x")
        v_scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        zoom_frame = tk.Frame(frame_left)
        zoom_frame.pack(pady=5)
        

        self.btn_zoom_in = tk.Button(zoom_frame, text="+", width=3, command=self.zoom_in)
        self.btn_zoom_in.pack(side="left", padx=5)

        self.btn_zoom_out = tk.Button(zoom_frame, text="-", width=3, command=self.zoom_out)
        self.btn_zoom_out.pack(side="left", padx=5)

        try:
            self.img = self.pdf_to_image(pdf_path)
            self.tk_img_ref = ImageTk.PhotoImage(self.img)
            self.image_on_canvas = self.canvas.create_image(0, 0, anchor="nw", image=self.tk_img_ref)
            self.canvas.config(scrollregion=self.canvas.bbox(tk.ALL))
        except Exception as e:
            tk.Label(frame_left, text="Не удалось отобразить PDF", font=("Arial", 12)).pack(padx=10, pady=10)
            print(f"[!] Ошибка при отображении PDF: {e}")

        # --- Выбор области ---
        self.start_x = self.start_y = 0
        self.end_x = self.end_y = 0
        self.rect = None

        self.canvas.bind("<ButtonPress-1>", self.on_mouse_down)
        self.canvas.bind("<B1-Motion>", self.on_mouse_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_mouse_up)

        # --- Правая часть: форма выбора данных ---
        frame_right = tk.Frame(root)
        frame_right.pack(side="right", fill="both", expand=False)

        # Поле: Материал (Combobox)
        tk.Label(frame_right, text="Материал:", font=("Arial", 14)).pack(anchor=tk.W, padx=20, pady=(10, 0))

        self.material_values = load_materials_from_excel(excel_file, sheet_name="Лист")  # Подгружаем материалы
        self.material_var = tk.StringVar(value=self.материал or "")

        self.material_combo = ttk.Combobox(frame_right, textvariable=self.material_var,
                                          values=self.material_values, font=("Arial", 12), width=38)
        self.material_combo.pack(padx=20, pady=5)
        self.material_combo.set(self.материал or "")

        # Тип заготовки (радиокнопки)
        self.radio_var = tk.StringVar()

        tk.Label(frame_right, text="Выберите тип заготовки:", font=("Arial", 14)).pack(anchor=tk.W, padx=20, pady=(10, 0))

        options_frame = tk.Frame(frame_right)
        options_frame.pack(anchor=tk.W, padx=20, pady=5)

        if os.path.exists(excel_file):
            self.types = load_types_from_excel(excel_file)
        else:
            self.types = {"Лист": True, "Круг": False, "Плита": False, "Сборка": False}

        for name, checked in self.types.items():
            rb = tk.Radiobutton(options_frame, text=name, value=name,
                               variable=self.radio_var, command=self.update_checkboxes)
            rb.pack(anchor=tk.W)
            if checked:
                rb.select()
                self.выбранный_тип = name
            else:
                rb.deselect()

        # Чекбоксы операций
        self.checkbox_frame = tk.Frame(frame_right)
        self.checkbox_frame.pack(anchor=tk.W, padx=20, pady=10)

        # Кнопки
        btn_frame = tk.Frame(frame_right)
        btn_frame.pack(side="bottom", pady=20)

        self.btn_continue = tk.Button(btn_frame, text="Продолжить", state=tk.DISABLED,
                                     command=self.on_continue, font=("Arial", 12), width=15, bg="lightblue")
        self.btn_continue.pack(side="left", padx=10)

        self.btn_skip = tk.Button(btn_frame, text="Пропустить", command=self.on_skip,
                                 font=("Arial", 12), width=15, bg="lightgray")
        self.btn_skip.pack(side="left", padx=10)

        # Инициализация чекбоксов под первый тип
        self.check_boxes = []
        self.check_vars = {}
        self.update_checkboxes()

    def pdf_to_image(self, pdf_path, zoom=None):
        """Конвертирует первую страницу PDF в PIL-изображение и масштабирует"""
        zoom = zoom or self.zoom_level
        doc = fitz.open(pdf_path)
        page = doc.load_page(0)

        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        doc.close()

        return Image.open(io.BytesIO(pix.tobytes("png")))

    def update_preview(self):
        try:
            self.img = self.pdf_to_image(self.pdf_path, self.zoom_level)
            self.tk_img_ref = ImageTk.PhotoImage(self.img)
            self.canvas.delete("all")
            self.image_on_canvas = self.canvas.create_image(0, 0, anchor="nw", image=self.tk_img_ref)
            self.canvas.config(scrollregion=self.canvas.bbox(tk.ALL))
        except Exception as e:
            print(f"[!] Не удалось обновить превью: {e}")

    def zoom_in(self):
        self.zoom_level *= 1.2
        self.update_preview()

    def zoom_out(self):
        self.zoom_level /= 1.2
        self.update_preview()

    def on_mouse_down(self, event):
        self.start_x = self.canvas.canvasx(event.x)
        self.start_y = self.canvas.canvasy(event.y)
        self.rect = self.canvas.create_rectangle(self.start_x, self.start_y, self.start_x, self.start_y, outline="red", width=2)

    def on_mouse_drag(self, event):
        cur_x = self.canvas.canvasx(event.x)
        cur_y = self.canvas.canvasy(event.y)
        self.canvas.coords(self.rect, self.start_x, self.start_y, cur_x, cur_y)

    def on_mouse_up(self, event):
        end_x = self.canvas.canvasx(event.x)
        end_y = self.canvas.canvasy(event.y)
        x1 = min(self.start_x, end_x)
        y1 = min(self.start_y, end_y)
        x2 = max(self.start_x, end_x)
        y2 = max(self.start_y, end_y)
        self.selected_image_area = (x1, y1, x2, y2)
        self.canvas.delete(self.rect)

    def update_checkboxes(self):
        selected_type = self.radio_var.get()
        if not selected_type:
            selected_type = next(iter(self.types), None)

        self.выбранный_тип = selected_type

        # Удаляем старые флажки
        for cb in self.check_boxes:
            cb.destroy()
        self.check_boxes.clear()
        self.check_vars.clear()

        # Обновляем материалы
        self.material_values = load_materials_from_excel(excel_file, sheet_name=selected_type)
        self.material_combo["values"] = self.material_values
        if self.material_values:
            self.material_combo.set(self.material_values[0])
        elif self.материал:
            self.material_combo.set(self.материал)

        # Загружаем операции
        operations = load_operations_with_columns(excel_file, selected_type)

        for op_name, is_checked in operations.items():
            if op_name == "nan":
                continue  # Пропускаем пустые/неправильные названия
            var = tk.BooleanVar(value=is_checked)
            cb = tk.Checkbutton(self.checkbox_frame, text=op_name, variable=var, command=self.toggle_continue)
            cb.pack(anchor=tk.W)
            self.check_vars[op_name] = var
            self.check_boxes.append(cb)

        self.toggle_continue()

    def toggle_continue(self):
        if any(var.get() for var in self.check_vars.values()):
            self.btn_continue.config(state=tk.NORMAL)
        else:
            self.btn_continue.config(state=tk.DISABLED)

    def on_continue(self):
        self.выбранные_операции = [op for op, var in self.check_vars.items() if var.get()]
        self.материал = self.material_var.get().strip()

        if not self.выбранный_тип:
            messagebox.showwarning("Ошибка", "Тип заготовки не выбран!")
            return

        if not self.выбранные_операции:
            messagebox.showwarning("Ошибка", "Выберите хотя бы одну операцию!")
            return

        self.root.quit()
        self.root.destroy()

    def on_skip(self):
        self.выбранный_тип = None
        self.выбранные_операции = []
        self.root.quit()
        self.root.destroy()


def replace_all(document, old_text, new_text):
    replaced = False
    old_text_cleaned = re.sub(r'\s+', '', old_text).lower()

    for para in document.paragraphs:
        full_text = ''.join(run.text for run in para.runs)
        if re.sub(r'\s+', '', full_text).lower() == old_text_cleaned:
            if para.runs:
                first_run = para.runs[0]
                bold = first_run.bold
                italic = first_run.italic
                underline = first_run.underline
                font_name = first_run.font.name
                font_size = first_run.font.size
                color = first_run.font.color.rgb
                hidden = first_run.font.hidden
            else:
                bold = italic = underline = font_name = font_size = color = hidden = None

            para.text = ""
            run = para.add_run(str(new_text or ""))
            run.bold = bold
            run.italic = italic
            run.underline = underline
            run.font.name = font_name
            run.font.size = font_size
            run.font.color.rgb = color
            run.font.hidden = hidden
            replaced = True

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = ''.join(run.text for run in para.runs)
                    if re.sub(r'\s+', '', full_text).lower() == old_text_cleaned:
                        if para.runs:
                            first_run = para.runs[0]
                            bold = first_run.bold
                            italic = first_run.italic
                            underline = first_run.underline
                            font_name = first_run.font.name
                            font_size = first_run.font.size
                            color = first_run.font.color.rgb
                            hidden = first_run.font.hidden
                        else:
                            bold = italic = underline = font_name = font_size = color = hidden = None

                        para.text = ""
                        run = para.add_run(str(new_text or ""))
                        run.bold = bold
                        run.italic = italic
                        run.underline = underline
                        run.font.name = font_name
                        run.font.size = font_size
                        run.font.color.rgb = color
                        run.font.hidden = hidden
                        replaced = True

    return replaced


def insert_image(document, keyword, image_path, width_inches=6):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    keyword_cleaned = re.sub(r'\s+', '', keyword).lower()

    for para in document.paragraphs:
        full_text = ''.join(run.text for run in para.runs)
        if re.sub(r'\s+', '', full_text).lower() == keyword_cleaned:
            para.text = ""
            run = para.add_run()
            run.add_picture(image_path, width=Inches(width_inches))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = ''.join(run.text for run in para.runs)
                    if re.sub(r'\s+', '', full_text).lower() == keyword_cleaned:
                        para.text = ""
                        run = para.add_run()
                        run.add_picture(image_path, width=Inches(width_inches))
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        return True

    print(f"[!] Метка '{keyword}' не найдена в документе")
    return False


def extract_and_save_image(pdf_path, output_path, area=None, zoom_preview=3, zoom_export=6):
    try:
        doc = fitz.open(pdf_path)
        page = doc.load_page(0)

        if area is None:
            clip_rect = page.rect
        else:
            orig_x1 = area[0] / zoom_preview
            orig_y1 = area[1] / zoom_preview
            orig_x2 = area[2] / zoom_preview
            orig_y2 = area[3] / zoom_preview
            clip_rect = fitz.Rect(orig_x1, orig_y1, orig_x2, orig_y2)

        mat_export = fitz.Matrix(zoom_export, zoom_export)
        pix_export = page.get_pixmap(matrix=mat_export, clip=clip_rect)
        pix_export.save(output_path, "png")
        doc.close()
        print(f"[OK] Изображение сохранено: {output_path}")
        return True
    except Exception as e:
        print(f"[!] Не удалось сохранить изображение: {e}")
        return False


def load_types_from_excel(excel_file):
    try:
        df = pd.read_excel(excel_file, sheet_name=0)
        types = {}
        for _, row in df.iterrows():
            name = str(row.iloc[0]).strip()
            state = str(row.iloc[1]).strip().lower()
            if name and name != "nan":
                types[name] = state == "checked"
        return types
    except Exception as e:
        print(f"[!] Ошибка чтения типов заготовки из Excel: {e}")
        return {"Лист": True}


def load_materials_from_excel(excel_file, sheet_name):
    try:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        materials = df["Материал"].astype(str).str.strip().dropna().unique().tolist()
        return materials
    except Exception as e:
        print(f"[!] Ошибка чтения материалов из Excel: {e}")
        return []


def get_operation_data(excel_file, sheet_name, operation_name):
    try:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        row = df[df.iloc[:, 0].astype(str).str.strip() == operation_name.strip()]
        if not row.empty:
            return {
                "description": str(row["ОписаниеТехноОперации"].values[0]).strip(),
                "equipment": str(row["ОборудованиеОперации"].values[0]).strip(),
                "tool": str(row["ПриспособлениеОперации"].values[0]).strip(),
                "storage": str(row["СкладОперации"].values[0]).strip()
            }
        else:
            print(f"[!] Не найдены данные для операции '{operation_name}' на листе '{sheet_name}'")
            return {"description": "", "equipment": "", "tool": "", "storage": ""}
    except Exception as e:
        print(f"[!] Ошибка получения данных операции: {e}")
        return {"description": "", "equipment": "", "tool": "", "storage": ""}


def load_operations_with_columns(excel_file, sheet_name):
    try:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        operations = {}
        for _, row in df.iterrows():
            op_name = str(row.iloc[0]).strip()
            state = str(row.iloc[1]).strip().lower()
            if op_name and op_name != "nan":
                operations[op_name] = state == "checked"
        return operations
    except Exception as e:
        print(f"[!] Ошибка чтения операций из листа '{sheet_name}': {e}")
        return {}


# --- Основной цикл ---
try:
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(".pdf"):
            file_path = os.path.join(folder_path, filename)

            try:
                doc_pdf = fitz.open(file_path)
                page = doc_pdf.load_page(0)

                is_large_sheet = page.rect.width > 800 or page.rect.height > 1000
                regions = rects_A3 if is_large_sheet else rects_A4

                ОбозначениеДетали = extract_text_from_rect(page, regions[0])
                НаименованиеДетали = extract_text_from_rect(page, regions[1])
                МатериалДетали = extract_text_from_rect(page, regions[2])
                РазработалДеталь = extract_text_from_rect(page, regions[3])
                ДатаЧертежа = extract_text_from_rect(page, regions[4])

                doc_pdf.saveIncr()
                doc_pdf.close()

                ОбозначениеДетали = normalize_dashes_and_spaces(ОбозначениеДетали)
                НаименованиеДетали = clean_filename(НаименованиеДетали)
                МатериалДетали = clean_filename(МатериалДетали)
                РазработалДеталь = clean_filename(РазработалДеталь)
                ДатаЧертежа = clean_filename(ДатаЧертежа)

                print(f"[DEBUG] {filename}:")
                print(f"Обозначение: {ОбозначениеДетали}")
                print(f"Наименование: {НаименованиеДетали}")
                print(f"Материал: {МатериалДетали}")
                print(f"Разработал: {РазработалДеталь}")
                print(f"Дата: {ДатаЧертежа}")

                # 🖼 Открываем GUI
                root = tk.Tk()
                app = App(root, file_path, ОбозначениеДетали, НаименованиеДетали, МатериалДетали, РазработалДеталь, ДатаЧертежа)

                try:
                    root.mainloop()
                except Exception as e:
                    print(f"[!] Ошибка GUI: {e}")
                    continue

                if not app.выбранные_операции and not app.выбранный_тип:
                    print(f"[SKIP] {filename} → пропущен.")
                    continue

                # 📄 Создание .docx
                base_name = os.path.splitext(filename)[0]
                docx_file = os.path.join(folder_path, f"{base_name}.docx")

                if not os.path.exists(template_docx_path):
                    print(f"[!] Шаблон '{template_docx_path}' не найден → пропуск файла")
                    continue

                try:
                    document = Document(template_docx_path)

                    # Замена базовых полей (все вхождения)
                    replace_all(document, "#ОбозначениеДетали", app.обозначение)
                    replace_all(document, "#НаименованиеДетали", app.наименование)
                    replace_all(document, "#МатериалДетали", app.material_var.get())
                    replace_all(document, "#РазработалДеталь", app.разработал)
                    replace_all(document, "#ДатаЧертежа", app.дата)
                    replace_all(document, "#ТипЗаготовки", app.выбранный_тип or "")

                    # Первая операция зависит от типа заготовки
                    if app.выбранный_тип != "Сборка":
                        replace_all(document, "#ОписаниеТехноОперации1",
                                    "Получение материала со склада в соответствии с сменным заданием и чертежом.")
                        replace_all(document, "#СкладОперации1", "Склад заготовительного участка механического цеха.")
                    else:
                        replace_all(document, "#ОписаниеТехноОперации1",
                                    "Получение деталей со склада в соответствии со спецификацией.")
                        replace_all(document, "#СкладОперации1", "Склад деталей сборочного цеха.")

                    replace_all(document, "#ОборудованиеОперации1", "")
                    replace_all(document, "#ПриспособлениеОперации1", "")

                    # Последующие операции
                    for idx, op in enumerate(app.выбранные_операции, start=2):
                        data = get_operation_data(excel_file, app.выбранный_тип, op)
                        replace_all(document, f"#ОписаниеТехноОперации{idx}", data["description"])
                        replace_all(document, f"#ОборудованиеОперации{idx}", data["equipment"])
                        replace_all(document, f"#ПриспособлениеОперации{idx}", data["tool"])
                        replace_all(document, f"#СкладОперации{idx}", data["storage"])

                    # Сохраняем изображение
                    temp_image_path = os.path.join(folder_path, "temp_selected.png")
                    area_to_use = app.selected_image_area if hasattr(app, 'selected_image_area') and app.selected_image_area else (
                        0, 0, app.tk_img_ref.width(), app.tk_img_ref.height()
                    )

                    if extract_and_save_image(file_path, temp_image_path, area_to_use, zoom_preview=app.zoom_level, zoom_export=6):
                        if not insert_image(document, "#Картинка", temp_image_path):
                            print("[!] Не удалось вставить картинку")
                    else:
                        print("[!] Не удалось сохранить выделенную область")

                    # Сохранение документа
                    document.save(docx_file)
                    print(f"[OK] Сохранён файл: {docx_file}")

                except Exception as e:
                    print(f"[Ошибка при работе с Word]: {e}")
                    continue
            except Exception as e:
                print(f"[Ошибка] Не удалось обработать '{filename}': {e}")
                try:
                    doc_pdf.close()
                except:
                    pass
                continue
except Exception as e:
    print(f"[Ошибка] Не удалось обработать файл: {e}")

messagebox.showinfo("Готово", "Обработка завершена.")