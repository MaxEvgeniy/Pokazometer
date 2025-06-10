import fitz  # PyMuPDF
import os
import re
from docx import Document
import tkinter as tk
from tkinter import messagebox
from PIL import Image, ImageTk
import io


# Путь к папке с PDF-файлами
folder_path = r'C:\PDF\1\Rename'

# Базовые регионы для формата A4
rects_A4 = [
    fitz.Rect(245, 676, 576, 709),   # Обозначение детали
    fitz.Rect(245, 720, 434, 750),   # Наименование детали
    fitz.Rect(245, 788, 434, 825),  # Материал детали
    fitz.Rect(105, 744, 170, 757),   # Разработал
    fitz.Rect(212, 744, 240, 757)    # Дата чертежа
]

# Коэффициенты смещения для A3
offset_x = 585
offset_y = 1
rects_A3 = [fitz.Rect(r.x0 + offset_x, r.y0 + offset_y, r.x1 + offset_x, r.y1 + offset_y) for r in rects_A4]

def normalize_dashes_and_spaces(text):
    """Замена всех видов дефисов/тире на '-'"""
    return re.sub(r'[\u00AD\u2010\u2011\u2012\u2013\u2014\u2015]', '-', text).strip()

def extract_text_from_rect(page, rect):
    """Извлечение текста из региона. Если нет — пробел."""
    clip_text = page.get_text("text", clip=rect).strip()
    lines = [line.strip() for line in clip_text.splitlines() if line.strip()]
    return lines[0] if lines else " "

def clean_filename(name):
    """Очистка имени файла от недопустимых символов"""
    name = re.sub(r'\s+', ' ', name).strip()
    return re.sub(r'[\\/*?:"<>|]', '', name)

def replace_all(doc, old_text, new_text):
    """Заменяет все вхождения старого текста на новый"""

    def replace_in_runs(runs):
        full_text = ''.join(run.text for run in runs)
        if old_text not in full_text:
            return False

        combined = full_text.replace(old_text, new_text)
        runs[0].text = combined
        for run in runs[1:]:
            run.text = ''
        return True

    replaced = False
    for para in doc.paragraphs:
        if old_text in para.text:
            if replace_in_runs(para.runs):
                replaced = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        replace_in_runs(para.runs)

    return replaced

class App:
    def __init__(self, root, pdf_path, обозначение, наименование):
        self.root = root
        self.pdf_path = pdf_path
        self.обозначение = обозначение
        self.наименование = наименование
        self.выбранный_тип = None

        # Установка заголовка окна
        root.title(f"{self.обозначение} {self.наименование}")

        # Левая часть: изображение PDF
        frame_left = tk.Frame(root, width=400, height=600)
        frame_left.pack(side="left", fill="both", expand=True)

        try:
            # Конвертируем первую страницу PDF в изображение
            img = self.pdf_to_image(pdf_path)
            img.thumbnail((380, 580))
            self.img_tk = ImageTk.PhotoImage(img)
            label_img = tk.Label(frame_left, image=self.img_tk)
            label_img.image = self.img_tk
            label_img.pack(padx=10, pady=10)
        except Exception as e:
            tk.Label(frame_left, text="Не удалось загрузить PDF", font=("Arial", 12)).pack(padx=10, pady=10)

        # Правая часть: радиокнопки
        frame_right = tk.Frame(root, width=400, height=600)
        frame_right.pack(side="right", fill="both", expand=True)

        tk.Label(frame_right, text="Выберите тип детали:", font=("Arial", 14)).pack(pady=20)

        self.radio_var = tk.StringVar(value="")
        self.radio_buttons = []

        for option in ["Лист", "Круг", "Плита"]:
            rb = tk.Radiobutton(frame_right, text=option, variable=self.radio_var, value=option, font=("Arial", 12))
            rb.pack(anchor=tk.W, padx=20)
            self.radio_buttons.append(rb)

        # Кнопка "Продолжить"
        self.btn_continue = tk.Button(frame_right, text="Продолжить", command=self.on_continue, font=("Arial", 12), width=15, bg="lightblue")
        self.btn_continue.pack(side="bottom", pady=20)

        self.continued = False

    def on_continue(self):
        selected = self.radio_var.get()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите тип детали!")
        else:
            self.выбранный_тип = selected
            self.root.destroy()
            self.continued = True

    def pdf_to_image(self, pdf_path, zoom=2):
        """Конвертирует первую страницу PDF в изображение PIL"""
        doc = fitz.open(pdf_path)
        page = doc.load_page(0)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        doc.close()
        return Image.open(io.BytesIO(img_data))


# Перебор всех PDF-файлов в папке
for filename in os.listdir(folder_path):
    if filename.lower().endswith(".pdf"):
        file_path = os.path.join(folder_path, filename)

        try:
            doc_pdf = fitz.open(file_path)
            page = doc_pdf.load_page(0)

            # Определение формата листа
            page_size = page.rect
            width, height = page_size.width, page_size.height
            is_large_sheet = width > 800 or height > 1000
            regions = rects_A3 if is_large_sheet else rects_A4

            print(f"[INFO] Формат листа: {'A3' if is_large_sheet else 'A4'}")

            # Извлечение данных из 5 регионов
            ОбозначениеДетали = extract_text_from_rect(page, regions[0])
            НаименованиеДетали = extract_text_from_rect(page, regions[1])
            МатериалДетали = extract_text_from_rect(page, regions[2])
            РазработалДеталь = extract_text_from_rect(page, regions[3])
            ДатаЧертежа = extract_text_from_rect(page, regions[4])

            doc_pdf.saveIncr()
            doc_pdf.close()

            # Нормализация строк
            ОбозначениеДетали = normalize_dashes_and_spaces(ОбозначениеДетали)
            НаименованиеДетали = clean_filename(НаименованиеДетали)
            МатериалДетали = clean_filename(МатериалДетали)
            РазработалДеталь = clean_filename(РазработалДеталь)

            print(f"[DEBUG] {filename}:")
            print(f"Обозначение: {ОбозначениеДетали}")
            print(f"Наименование: {НаименованиеДетали}")
            print(f"Материал: {МатериалДетали}")
            print(f"Разработал: {РазработалДеталь}")
            print(f"Дата: {ДатаЧертежа}")

            # 🖼 Показываем GUI
            root = tk.Tk()
            app = App(root, file_path, ОбозначениеДетали, НаименованиеДетали)

            root.mainloop()  # Ждём, пока пользователь выберет тип

            # Получаем выбранный тип
            тип_детали = app.radio_var.get()
            if not тип_детали:
                print("[!] Тип детали не выбран → пропуск файла")
                continue

            print(f"[OK] Выбрано: {тип_детали}")

            # 📄 Создание .docx на основе шаблона
            template_docx_path = os.path.join(folder_path, "template.docx")
            if not os.path.exists(template_docx_path):
                print(f"[!] Шаблон '{template_docx_path}' не найден → пропуск создания .docx")
                continue

            base_name = os.path.splitext(filename)[0]
            docx_file = os.path.join(folder_path, f"{base_name}.docx")

            try:
                document = Document(template_docx_path)

                replace_all(document, "#ОбозначениеДетали", ОбозначениеДетали)
                replace_all(document, "#НаименованиеДетали", НаименованиеДетали)
                replace_all(document, "#МатериалДетали", МатериалДетали)
                replace_all(document, "#РазработалДеталь", РазработалДеталь)
                replace_all(document, "#ДатаЧертежа", ДатаЧертежа)
                replace_all(document, "#ТипДетали", тип_детали)

                # Сохраняем документ
                document.save(docx_file)
                print(f"[OK] Сохранён файл: {docx_file}")

            except Exception as e:
                print(f"[Ошибка при создании .docx]: {e}")

        except Exception as e:
            print(f"[Ошибка] {filename}: {e}")

input("\nНажмите Enter для выхода...")