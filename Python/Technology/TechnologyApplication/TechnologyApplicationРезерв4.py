import fitz  # PyMuPDF
import os
import re
from docx import Document
import tkinter as tk
from tkinter import messagebox
from PIL import Image, ImageTk
import io


# Путь к папке с PDF-файлами
folder_path = r'C:\PDF\1\Rename'

# Базовые регионы для формата A4
rects_A4 = [
    fitz.Rect(245, 676, 576, 709),   # Обозначение детали
    fitz.Rect(245, 720, 434, 750),   # Наименование детали
    fitz.Rect(245, 788, 434, 825),   # Материал детали
    fitz.Rect(105, 744, 170, 757),   # Разработал
    fitz.Rect(212, 744, 240, 757)    # Дата чертежа
]

# Коэффициенты смещения для A3
offset_x = 585
offset_y = 1
rects_A3 = [fitz.Rect(r.x0 + offset_x, r.y0 + offset_y, r.x1 + offset_x, r.y1 + offset_y) for r in rects_A4]


def normalize_dashes_and_spaces(text):
    return re.sub(r'[\u00AD\u2010\u2011\u2012\u2013\u2014\u2015]', '-', text).strip()

def extract_text_from_rect(page, rect):
    clip_text = page.get_text("text", clip=rect).strip()
    lines = [line.strip() for line in clip_text.splitlines() if line.strip()]
    return lines[0] if lines else " "

def clean_filename(name):
    name = re.sub(r'\s+', ' ', name).strip()
    return re.sub(r'[\\/*?:"<>|]', '', name)


class App:
    def __init__(self, root, pdf_path, обозначение, наименование, материал, разработал, дата):
        self.root = root
        self.pdf_path = pdf_path
        self.обозначение = обозначение
        self.наименование = наименование
        self.материал = материал
        self.разработал = разработал
        self.дата = дата
        self.выбранный_тип = None
        self.выбранные_операции = []

        # Установка заголовка окна
        root.title(f"{self.обозначение} {self.наименование}")
        root.geometry("800x600")

        # Левая часть: превью первой страницы PDF
        frame_left = tk.Frame(root, width=400, height=600)
        frame_left.pack(side="left", fill="both", expand=True)

        try:
            img = self.pdf_to_image(pdf_path)
            img.thumbnail((380, 580))
            self.img_tk = ImageTk.PhotoImage(img)
            label_img = tk.Label(frame_left, image=self.img_tk)
            label_img.image = self.img_tk
            label_img.pack(padx=10, pady=10)
        except Exception as e:
            tk.Label(frame_left, text="Не удалось загрузить PDF", font=("Arial", 12)).pack(padx=10, pady=10)

        # Правая часть: радиокнопки и чекбоксы
        frame_right = tk.Frame(root, width=400, height=600)
        frame_right.pack(side="right", fill="both", expand=True)

        # Выбор типа заготовки
        self.radio_var = tk.StringVar(value="Лист")
        tk.Label(frame_right, text="Выберите тип заготовки:", font=("Arial", 14)).pack(anchor=tk.W, padx=20, pady=(10, 0))

        self.options_frame = tk.Frame(frame_right)
        self.options_frame.pack(anchor=tk.W, padx=20, pady=10)

        self.radio_buttons = []
        for option in ["Лист", "Круг", "Плита"]:
            rb = tk.Radiobutton(self.options_frame, text=option, value=option,
                                variable=self.radio_var, command=self.update_checkboxes)
            rb.pack(anchor=tk.W)
            self.radio_buttons.append(rb)

        # Фрейм для чекбоксов
        self.checkbox_frame = tk.Frame(frame_right)
        self.checkbox_frame.pack(anchor=tk.W, padx=20, pady=10)

        # Словарь операций по типу заготовки
        self.operations = {
            "Лист": ["Вырубка заготовки", "Пробивка", "Гибка", "Нарезка резьбы"],
            "Круг": ["Распил заготовки", "Токарка", "Сверление", "Нарезка резьбы"],
            "Плита": ["Распил заготовки", "Фрезеровка", "Токарка", "Сверление", "Нарезка резьбы"]
        }

        # Значения по умолчанию
        self.default_checks = {
            "Лист": ["Вырубка заготовки", "Пробивка", "Гибка"],
            "Круг": ["Распил заготовки", "Токарка"],
            "Плита": ["Распил заготовки", "Фрезеровка"]
        }

        # Для хранения чекбоксов
        self.check_vars = {}
        self.check_boxes = []

        # Кнопка "Продолжить"
        self.btn_continue = tk.Button(frame_right, text="Продолжить", state=tk.DISABLED,
                                      command=self.on_continue, font=("Arial", 12), width=15, bg="lightblue")
        self.btn_continue.pack(side="bottom", pady=20)

        # Инициализация чекбоксов
        self.update_checkboxes()

    def pdf_to_image(self, pdf_path, zoom=2):
        """Конвертирует первую страницу PDF в изображение PIL"""
        doc = fitz.open(pdf_path)
        page = doc.load_page(0)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        doc.close()
        return Image.open(io.BytesIO(img_data))

    def update_checkboxes(self):
        """Обновляет список чекбоксов в зависимости от выбранного типа заготовки"""

        # Очистка старых флажков
        for cb in self.check_boxes:
            cb.destroy()
        self.check_boxes.clear()
        self.check_vars.clear()

        self.selected_type = self.radio_var.get()
        ops = self.operations[self.selected_type]

        # Создание новых флажков
        for op in ops:
            var = tk.BooleanVar(value=op in self.default_checks[self.selected_type])
            cb = tk.Checkbutton(self.checkbox_frame, text=op, variable=var, command=self.toggle_continue)
            cb.pack(anchor=tk.W)
            self.check_vars[op] = var
            self.check_boxes.append(cb)

        self.toggle_continue()

    def toggle_continue(self):
        """Активирует/деактивирует кнопку Продолжить"""
        if any(var.get() for var in self.check_vars.values()):
            self.btn_continue.config(state=tk.NORMAL)
        else:
            self.btn_continue.config(state=tk.DISABLED)

    def on_continue(self):
        self.выбранный_тип = self.radio_var.get()
        self.выбранные_операции = [op for op, var in self.check_vars.items() if var.get()]

        if not self.выбранный_тип:
            messagebox.showwarning("Ошибка", "Тип заготовки не выбран!")
            return

        if not self.выбранные_операции:
            messagebox.showwarning("Ошибка", "Выберите хотя бы одну операцию!")
            return

        self.root.destroy()
        print("[OK] Тип заготовки:", self.выбранный_тип)
        print("[OK] Операции:", ", ".join(self.выбранные_операции))


# --- Функция замены ключевых слов ---
def replace_all(document, old_text, new_text):
    """Заменяет все вхождения old_text на new_text, сохраняя структуру и стиль run'ов"""
    
    def replace_in_runs(runs, old_text, new_text):
        full_text = ''.join(run.text for run in runs)
        
        if old_text not in full_text:
            return False

        # Разбиваем весь текст на части до и после ключевого слова
        start_idx = full_text.find(old_text)
        end_idx = start_idx + len(old_text)

        before = full_text[:start_idx]
        after = full_text[end_idx:]

        # Сохраняем стиль первого run
        first_run = runs[0]
        bold = first_run.bold
        italic = first_run.italic
        underline = first_run.underline
        font_name = first_run.font.name
        font_size = first_run.font.size
        try:
            color = first_run.font.color.rgb
        except:
            color = None

        # Очищаем старые run'ы
        for run in runs:
            run.text = ''

        # Восстанавливаем текст с новым значением
        runs[0].text = before + new_text + after
        runs[0].bold = bold
        runs[0].italic = italic
        runs[0].underline = underline
        runs[0].font.name = font_name
        runs[0].font.size = font_size
        if color:
            runs[0].font.color.rgb = color

        return True

    replaced = False
    
    # Параграфы
    for para in document.paragraphs:
        if old_text in para.text:
            if replace_in_runs(para.runs, old_text, new_text):
                replaced = True

    # Таблицы
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        replace_in_runs(para.runs, old_text, new_text)

    return replaced


# --- Основной цикл ---
for filename in os.listdir(folder_path):
    if filename.lower().endswith(".pdf"):
        file_path = os.path.join(folder_path, filename)

        try:
            doc_pdf = fitz.open(file_path)
            page = doc_pdf.load_page(0)

            # Определение формата листа
            page_size = page.rect
            width, height = page_size.width, page_size.height
            is_large_sheet = width > 800 or height > 1000
            regions = rects_A3 if is_large_sheet else rects_A4

            print(f"[INFO] Формат листа: {'A3' if is_large_sheet else 'A4'}")

            # Извлечение данных из 5 регионов
            ОбозначениеДетали = extract_text_from_rect(page, regions[0])
            НаименованиеДетали = extract_text_from_rect(page, regions[1])
            МатериалДетали = extract_text_from_rect(page, regions[2])
            РазработалДеталь = extract_text_from_rect(page, regions[3])
            ДатаЧертежа = extract_text_from_rect(page, regions[4])

            doc_pdf.saveIncr()
            doc_pdf.close()

            # Нормализация строк
            ОбозначениеДетали = normalize_dashes_and_spaces(ОбозначениеДетали)
            НаименованиеДетали = clean_filename(НаименованиеДетали)
            МатериалДетали = clean_filename(МатериалДетали)
            РазработалДеталь = clean_filename(РазработалДеталь)
            ДатаЧертежа = clean_filename(ДатаЧертежа)

            print(f"[DEBUG] {filename}:")
            print(f"Обозначение: {ОбозначениеДетали}")
            print(f"Наименование: {НаименованиеДетали}")
            print(f"Материал: {МатериалДетали}")
            print(f"Разработал: {РазработалДеталь}")
            print(f"Дата: {ДатаЧертежа}")

            # 🖼 Показываем GUI
            root = tk.Tk()
            app = App(root, file_path, ОбозначениеДетали, НаименованиеДетали, МатериалДетали, РазработалДеталь, ДатаЧертежа)

            try:
                root.mainloop()
            except Exception as e:
                print(f"[!] Ошибка при показе окна: {e}")
                continue

            # Получаем результаты после закрытия окна
            тип_заготовки = app.выбранный_тип
            операции = app.выбранные_операции
            операции_str = "\n".join(операции)

            print(f"[OK] Выбрано: {тип_заготовки}, Операции: {операции_str}")

            # 📄 Создание .docx на основе шаблона
            template_docx_path = os.path.join(folder_path, "template.docx")
            if not os.path.exists(template_docx_path):
                print(f"[!] Шаблон '{template_docx_path}' не найден → пропуск создания .docx")
                continue

            base_name = os.path.splitext(filename)[0]
            docx_file = os.path.join(folder_path, f"{base_name}.docx")

            try:
                document = Document(template_docx_path)

                replace_all(document, "#ОбозначениеДетали", ОбозначениеДетали)
                replace_all(document, "#НаименованиеДетали", НаименованиеДетали)
                replace_all(document, "#МатериалДетали", МатериалДетали)
                replace_all(document, "#РазработалДеталь", РазработалДеталь)
                replace_all(document, "#ДатаЧертежа", ДатаЧертежа)
                replace_all(document, "#ТипЗаготовки", тип_заготовки)
                replace_all(document, "#Операции", операции_str)

                # Замена #ОписаниеТехноОперации по выбранному типу заготовки
                if тип_заготовки == "Лист":
                    описание_по_типу = "Получение материала со склада в соответствии с сменным заданием и чертежом."
                elif тип_заготовки == "Круг":
                    описание_по_типу = "Использование круга согласно чертежу."
                elif тип_заготовки == "Плита":
                    описание_по_типу = "Использование плиты согласно чертежу."
                else:
                    описание_по_типу = ""

                replace_all(document, "#ОписаниеТехноОперации", описание_по_типу)

                # Сохраняем документ
                document.save(docx_file)
                print(f"[OK] Сохранён файл: {docx_file}")

            except Exception as e:
                print(f"[Ошибка при работе с Word]: {e}")

        except Exception as e:
            print(f"[Ошибка] Не удалось обработать '{filename}': {e}")
            try:
                doc_pdf.close()
            except:
                pass
            continue

messagebox.showinfo("Готово", "Обработка завершена.")
input("\nНажмите Enter для выхода...")