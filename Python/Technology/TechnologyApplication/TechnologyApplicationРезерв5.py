import fitz  # PyMuPDF
import os
import re
from docx import Document
import tkinter as tk
from tkinter import messagebox
from docx.shared import Pt, RGBColor
from PIL import Image, ImageTk
import io


# Путь к папке с PDF-файлами
folder_path = r'C:\PDF\1\Rename'

# Базовые регионы для A4
rects_A4 = [
    fitz.Rect(245, 676, 576, 709),   # Обозначение детали
    fitz.Rect(245, 720, 434, 750),   # Наименование детали
    fitz.Rect(245, 788, 434, 825),   # Материал детали
    fitz.Rect(105, 744, 170, 757),   # Разработал
    fitz.Rect(212, 744, 240, 757)    # Дата чертежа
]

# Смещения для A3
offset_x = 585
offset_y = 1
rects_A3 = [fitz.Rect(r.x0 + offset_x, r.y0 + offset_y, r.x1 + offset_x, r.y1 + offset_y) for r in rects_A4]


def normalize_dashes_and_spaces(text):
    return re.sub(r'[\u00AD\u2010\u2011\u2012\u2013\u2014\u2015]', '-', text).strip()

def extract_text_from_rect(page, rect):
    clip_text = page.get_text("text", clip=rect).strip()
    lines = [line.strip() for line in clip_text.splitlines() if line.strip()]
    return lines[0] if lines else " "

def clean_filename(name):
    name = re.sub(r'\s+', ' ', name).strip()
    return re.sub(r'[\\/*?:"<>|]', '', name)


class App:
    def __init__(self, root, pdf_path, обозначение, наименование, материал, разработал, дата):
        self.root = root
        self.pdf_path = pdf_path
        self.обозначение = обозначение
        self.наименование = наименование
        self.материал = материал
        self.разработал = разработал
        self.дата = дата
        self.выбранный_тип = None
        self.выбранные_операции = []

        # --- ВСЕГДА существуют ---
        self.check_boxes = []
        self.check_vars = {}

        # Установка заголовка окна
        root.title(f"{self.обозначение} {self.наименование}")
        root.geometry("800x600")

        # Левая часть: превью первой страницы PDF
        frame_left = tk.Frame(root, width=400, height=600)
        frame_left.pack(side="left", fill="both", expand=True)

        self.img_tk = None

        try:
            img = self.pdf_to_image(pdf_path)
            if img:
                img.thumbnail((380, 580))
                self.img_tk = ImageTk.PhotoImage(img)
                label_img = tk.Label(frame_left, image=self.img_tk)
                label_img.image = self.img_tk
                label_img.pack(padx=10, pady=10)
            else:
                raise Exception("Изображение не было создано")
        except Exception as e:
            tk.Label(frame_left, text="Не удалось отобразить PDF", font=("Arial", 12)).pack(padx=10, pady=10)
            print(f"[!] Ошибка при отображении PDF: {e}")

        # Правая часть: радиокнопки и чекбоксы
        frame_right = tk.Frame(root, width=400, height=600)
        frame_right.pack(side="right", fill="both", expand=True)

        # Выбор типа заготовки
        self.radio_var = tk.StringVar(value="Лист")
        tk.Label(frame_right, text="Выберите тип заготовки:", font=("Arial", 14)).pack(anchor=tk.W, padx=20, pady=(10, 0))

        self.options_frame = tk.Frame(frame_right)
        self.options_frame.pack(anchor=tk.W, padx=20, pady=10)

        self.operations = {
            "Лист": ["Вырубка заготовки", "Пробивка", "Гибка", "Нарезка резьбы"],
            "Круг": ["Распил заготовки", "Токарка", "Сверление", "Нарезка резьбы"],
            "Плита": ["Распил заготовки", "Фрезеровка", "Токарка", "Сверление", "Нарезка резьбы"]
        }

        self.default_checks = {
            "Лист": ["Вырубка заготовки", "Пробивка", "Гибка"],
            "Круг": ["Распил заготовки", "Токарка"],
            "Плита": ["Распил заготовки", "Фрезеровка"]
        }

        for option in ["Лист", "Круг", "Плита"]:
            rb = tk.Radiobutton(self.options_frame, text=option, value=option,
                                variable=self.radio_var, command=self.update_checkboxes)
            rb.pack(anchor=tk.W)

        # Чекбоксы
        self.checkbox_frame = tk.Frame(frame_right)
        self.checkbox_frame.pack(anchor=tk.W, padx=20, pady=10)

        # Кнопки
        btn_frame = tk.Frame(frame_right)
        btn_frame.pack(side="bottom", pady=20)

        self.btn_continue = tk.Button(btn_frame, text="Продолжить", state=tk.DISABLED,
                                      command=self.on_continue, font=("Arial", 12), width=15, bg="lightblue")
        self.btn_continue.pack(side="left", padx=10)

        self.btn_skip = tk.Button(btn_frame, text="Пропустить", command=self.on_skip,
                                  font=("Arial", 12), width=15, bg="lightgray")
        self.btn_skip.pack(side="left", padx=10)

        # Словарь описаний операций
        self.operation_descriptions = {
            "Вырубка заготовки": "Отделение заготовки по чертежу с припуском на дальнейшую обработку.",
            "Пробивка": "Пробивка отверстий в заготовке в соответствии с чертежом.",
            "Гибка": "Гибка заготовки в соответствии с чертежом.",
            "Распил заготовки": "Отделение заготовки по чертежу с припуском на дальнейшую обработку.",
            "Токарка": "Выполнение токарных операций в заготовке по чертежу.",
            "Сверление": "Выполнение отверстий в заготовке по чертежу.",
            "Нарезка резьбы": "Выполнение резьбы по чертежу.",
            "Фрезеровка": "Выполнение фрезеровальных операций в заготовке по чертежу."
        }

        # Инициализация чекбоксов
        self.update_checkboxes()

    def pdf_to_image(self, pdf_path, zoom=2):
        """Конвертирует первую страницу PDF в изображение PIL"""
        try:
            doc = fitz.open(pdf_path)
            page = doc.load_page(0)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            doc.close()
            return Image.open(io.BytesIO(img_data))
        except Exception as e:
            print(f"[!] Не удалось создать превью PDF: {e}")
            return None

    def update_checkboxes(self):
        """Обновляет список чекбоксов в зависимости от выбранного типа заготовки"""

        # Очистка старых флажков
        for cb in self.check_boxes:
            cb.destroy()
        self.check_boxes.clear()
        self.check_vars.clear()

        selected_type = self.radio_var.get()
        ops = self.operations[selected_type]

        # Создание новых флажков
        for op in ops:
            var = tk.BooleanVar(value=op in self.default_checks[selected_type])
            cb = tk.Checkbutton(self.checkbox_frame, text=op, variable=var, command=self.toggle_continue)
            cb.pack(anchor=tk.W)
            self.check_vars[op] = var
            self.check_boxes.append(cb)

        self.toggle_continue()

    def toggle_continue(self):
        """Активирует/деактивирует кнопку Продолжить"""
        if any(var.get() for var in self.check_vars.values()):
            self.btn_continue.config(state=tk.NORMAL)
        else:
            self.btn_continue.config(state=tk.DISABLED)

    def on_continue(self):
        self.выбранный_тип = self.radio_var.get()
        self.выбранные_операции = [op for op, var in self.check_vars.items() if var.get()]

        if not self.выбранный_тип:
            messagebox.showwarning("Ошибка", "Тип заготовки не выбран!")
            return

        if not self.выбранные_операции:
            messagebox.showwarning("Ошибка", "Выберите хотя бы одну операцию!")
            return

        self.root.quit()
        self.root.destroy()

    def on_skip(self):
        """Пропускает текущий файл"""
        print("[SKIP] Файл пропущен.")
        self.выбранный_тип = None
        self.выбранные_операции = []
        self.root.quit()
        self.root.destroy()


# --- Функция замены текста ---
def replace_all(document, old_text, new_text):
    replaced = False

    # Параграфы
    for para in document.paragraphs:
        if old_text in para.text:
            inline = para.runs
            full_text = ''.join(run.text for run in inline)
            combined = full_text.replace(old_text, new_text or "")

            para.clear()
            new_run = para.add_run(combined)
            if inline:
                src_run = inline[0]
                new_run.bold = src_run.bold
                new_run.italic = src_run.italic
                new_run.underline = src_run.underline
                try:
                    new_run.font.name = src_run.font.name
                    new_run.font.size = src_run.font.size
                except:
                    pass
            replaced = True

    # Таблицы
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        inline = para.runs
                        full_text = ''.join(run.text for run in inline)
                        combined = full_text.replace(old_text, new_text or "")
                        para.clear()
                        new_run = para.add_run(combined)
                        if inline:
                            src_run = inline[0]
                            new_run.bold = src_run.bold
                            new_run.italic = src_run.italic
                            new_run.underline = src_run.underline
                            try:
                                new_run.font.name = src_run.font.name
                                new_run.font.size = src_run.font.size
                            except:
                                pass
                        replaced = True

    return replaced


# --- Функция удаления текста между маркерами ---
def remove_text_between(document, start_marker="#Удалять", end_marker="#НЕУдалять"):
    """
    Удаляет весь текст между start_marker и end_marker (включая их).
    Также удаляет маркеры из текста.
    """

    deleting = False

    # Проходим по параграфам
    paras_to_clear = []
    for para in document.paragraphs:
        if start_marker in para.text:
            deleting = True
            para.text = para.text.replace(start_marker, "").strip()
            paras_to_clear.append(para)
        elif end_marker in para.text:
            deleting = False
            para.text = para.text.replace(end_marker, "").strip()
            paras_to_clear.append(para)
        elif deleting:
            para.text = ""
            paras_to_clear.append(para)

    # Теперь проходим по ячейкам таблиц
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if start_marker in para.text:
                        deleting = True
                        para.text = para.text.replace(start_marker, "").strip()
                    elif end_marker in para.text:
                        deleting = False
                        para.text = para.text.replace(end_marker, "").strip()
                    elif deleting:
                        para.text = ""

    print(f"[OK] Удалено {len(paras_to_clear)} параграфов между '{start_marker}' и '{end_marker}'.")

def remove_all_hash_tags(document):
    """
    Удаляет все вхождения слов, начинающихся с '#', из всего документа,
    включая параграфы и таблицы.
    """

    # Регулярное выражение для поиска #... 
    pattern = re.compile(r'#\S+')

    # --- Параграфы ---
    for para in document.paragraphs:
        if '#' in para.text:
            para.text = pattern.sub('', para.text).strip()

    # --- Таблицы ---
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '#' in para.text:
                        para.text = pattern.sub('', para.text).strip()

    print("[OK] Все ключевые слова вида #... удалены.")

# --- Основной цикл ---
for filename in os.listdir(folder_path):
    if filename.lower().endswith(".pdf"):
        file_path = os.path.join(folder_path, filename)

        try:
            doc_pdf = fitz.open(file_path)
            page = doc_pdf.load_page(0)

            # Определение формата листа
            page_size = page.rect
            width, height = page_size.width, page_size.height
            is_large_sheet = width > 800 or height > 1000
            regions = rects_A3 if is_large_sheet else rects_A4

            print(f"[INFO] Формат листа: {'A3' if is_large_sheet else 'A4'}")

            # Извлечение данных из регионов
            ОбозначениеДетали = extract_text_from_rect(page, regions[0])
            НаименованиеДетали = extract_text_from_rect(page, regions[1])
            МатериалДетали = extract_text_from_rect(page, regions[2])
            РазработалДеталь = extract_text_from_rect(page, regions[3])
            ДатаЧертежа = extract_text_from_rect(page, regions[4])

            doc_pdf.saveIncr()
            doc_pdf.close()

            # Нормализация строк
            ОбозначениеДетали = normalize_dashes_and_spaces(ОбозначениеДетали)
            НаименованиеДетали = clean_filename(НаименованиеДетали)
            МатериалДетали = clean_filename(МатериалДетали)
            РазработалДеталь = clean_filename(РазработалДеталь)
            ДатаЧертежа = clean_filename(ДатаЧертежа)

            print(f"[DEBUG] {filename}:")
            print(f"Обозначение: {ОбозначениеДетали}")
            print(f"Наименование: {НаименованиеДетали}")
            print(f"Материал: {МатериалДетали}")
            print(f"Разработал: {РазработалДеталь}")
            print(f"Дата: {ДатаЧертежа}")

            # 🖼 Открываем GUI
            root = tk.Tk()
            app = App(root, file_path, ОбозначениеДетали, НаименованиеДетали, МатериалДетали, РазработалДеталь, ДатаЧертежа)

            try:
                root.mainloop()
            except Exception as e:
                print(f"[!] Ошибка GUI: {e}")
                continue

            # Если пользователь нажал "Пропустить"
            if not app.выбранные_операции and not app.выбранный_тип:
                print(f"[SKIP] {filename} → пропущен.")
                continue

            # 📄 Создание .docx
            template_docx_path = os.path.join(folder_path, "template.docx")
            if not os.path.exists(template_docx_path):
                print(f"[!] Шаблон '{template_docx_path}' не найден → пропуск файла")
                continue

            base_name = os.path.splitext(filename)[0]
            docx_file = os.path.join(folder_path, f"{base_name}.docx")

            try:
                document = Document(template_docx_path)

                # Замена базовых полей
                replace_all(document, "#ОбозначениеДетали", ОбозначениеДетали or "")
                replace_all(document, "#НаименованиеДетали", НаименованиеДетали or "")
                replace_all(document, "#МатериалДетали", МатериалДетали or "")
                replace_all(document, "#РазработалДеталь", РазработалДеталь or "")
                replace_all(document, "#ДатаЧертежа", ДатаЧертежа or "")
                replace_all(document, "#ТипЗаготовки", app.выбранный_тип or "")

                # Первый лист — получение материала
                replace_all(document, "#ОписаниеТехноОперации1",
                            "Получение материала со склада в соответствии с сменным заданием и чертежом.")

                # Последующие листы — операции
                for idx, op in enumerate(app.выбранные_операции, start=2):
                    desc = app.operation_descriptions.get(op, f"Описание операции для {op} отсутствует.")
                    replace_all(document, f"#ОписаниеТехноОперации{idx}", desc or "")

                # Удаляем всё между маркерами
                remove_text_between(document)

                # Удалить всё, что осталось с #
                remove_all_hash_tags(document)

                # Сохраняем документ
                document.save(docx_file)
                print(f"[OK] Сохранён файл: {docx_file}")

            except Exception as e:
                print(f"[Ошибка при работе с Word]: {e}")
                continue

        except Exception as e:
            print(f"[Ошибка] Не удалось обработать '{filename}': {e}")
            try:
                doc_pdf.close()
            except:
                pass
            continue

messagebox.showinfo("Готово", "Обработка завершена.")
input("\nНажмите Enter для выхода...")