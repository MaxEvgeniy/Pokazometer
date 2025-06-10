import fitz  # PyMuPDF
import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import messagebox
from PIL import Image, ImageTk
import io


# Путь к папке с PDF-файлами
folder_path = r'C:\PDF\1\Rename'

# Регионы A4
rects_A4 = [
    fitz.Rect(245, 676, 576, 709),   # Обозначение детали
    fitz.Rect(245, 720, 434, 750),   # Наименование детали
    fitz.Rect(245, 788, 434, 825),   # Материал детали
    fitz.Rect(105, 744, 170, 757),   # Разработал
    fitz.Rect(212, 744, 240, 757)     # Дата чертежа
]

# Смещения для A3
offset_x = 585
offset_y = 1
rects_A3 = [fitz.Rect(r.x0 + offset_x, r.y0 + offset_y, r.x1 + offset_x, r.y1 + offset_y) for r in rects_A4]


def normalize_dashes_and_spaces(text):
    return re.sub(r'[\u00AD\u2010\u2011\u2012\u2013\u2014\u2015]', '-', text).strip()

def extract_text_from_rect(page, rect):
    clip_text = page.get_text("text", clip=rect).strip()
    lines = [line.strip() for line in clip_text.splitlines() if line.strip()]
    return lines[0] if lines else ""

def clean_filename(name):
    name = re.sub(r'\s+', ' ', name).strip()
    return re.sub(r'[\\/*?:"<>|]', '', name)


class App:
    def __init__(self, root, pdf_path, обозначение, наименование, материал, разработал, дата):
        self.root = root
        self.pdf_path = pdf_path
        self.обозначение = обозначение
        self.наименование = наименование
        self.материал = материал
        self.разработал = разработал
        self.дата = дата
        self.выбранный_тип = None
        self.выбранные_операции = []
        self.selected_image_area = None
        self.tk_img_ref = None
        self.zoom_level = 1  # начальный масштаб

        # --- GUI ---
        root.title(f"{self.обозначение} {self.наименование}")
        root.geometry("1200x800")

        # Левая часть: превью PDF с прокруткой
        frame_left = tk.Frame(root)
        frame_left.pack(side="left", fill="both", expand=True)

        canvas_frame = tk.Frame(frame_left)
        canvas_frame.pack(fill="both", expand=True)

        self.canvas = tk.Canvas(canvas_frame, bg="white")
        h_scrollbar = tk.Scrollbar(canvas_frame, orient="horizontal", command=self.canvas.xview)
        v_scrollbar = tk.Scrollbar(canvas_frame, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(xscrollcommand=h_scrollbar.set, yscrollcommand=v_scrollbar.set)

        h_scrollbar.pack(side="bottom", fill="x")
        v_scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        # Кнопки зума
        zoom_frame = tk.Frame(frame_left)
        zoom_frame.pack(pady=5)

        self.btn_zoom_in = tk.Button(zoom_frame, text="+", width=3, command=self.zoom_in)
        self.btn_zoom_in.pack(side="left", padx=5)

        self.btn_zoom_out = tk.Button(zoom_frame, text="-", width=3, command=self.zoom_out)
        self.btn_zoom_out.pack(side="left", padx=5)

        # Превью PDF
        try:
            self.img = self.pdf_to_image(pdf_path)
            self.tk_img_ref = ImageTk.PhotoImage(self.img)
            self.image_on_canvas = self.canvas.create_image(0, 0, anchor="nw", image=self.tk_img_ref)
            self.canvas.config(scrollregion=self.canvas.bbox(tk.ALL))
        except Exception as e:
            tk.Label(frame_left, text="Не удалось отобразить PDF", font=("Arial", 12)).pack(padx=10, pady=10)
            print(f"[!] Ошибка при отображении PDF: {e}")

        # --- Выбор области ---
        self.start_x = self.start_y = 0
        self.end_x = self.end_y = 0
        self.rect = None

        self.canvas.bind("<ButtonPress-1>", self.on_mouse_down)
        self.canvas.bind("<B1-Motion>", self.on_mouse_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_mouse_up)

        # --- Правая часть: форма выбора данных ---
        frame_right = tk.Frame(root)
        frame_right.pack(side="right", fill="both", expand=False)

        # Поле: Материал
        tk.Label(frame_right, text="Материал:", font=("Arial", 14)).pack(anchor=tk.W, padx=20, pady=(10, 0))
        self.material_entry = tk.Entry(frame_right, width=40, font=("Arial", 12))
        self.material_entry.insert(0, self.материал or "")
        self.material_entry.pack(padx=20, pady=5)

        # Выбор типа заготовки
        self.radio_var = tk.StringVar(value="Лист")
        tk.Label(frame_right, text="Выберите тип заготовки:", font=("Arial", 14)).pack(anchor=tk.W, padx=20, pady=(10, 0))

        options_frame = tk.Frame(frame_right)
        options_frame.pack(anchor=tk.W, padx=20, pady=5)

        self.operations = {
            "Лист": ["Вырубка заготовки", "Пробивка", "Гибка", "Нарезка резьбы"],
            "Круг": ["Распил заготовки", "Токарка", "Сверление", "Нарезка резьбы"],
            "Плита": ["Распил заготовки", "Фрезеровка", "Токарка", "Сверление", "Нарезка резьбы"]
        }

        self.default_checks = {
            "Лист": ["Вырубка заготовки", "Пробивка", "Гибка"],
            "Круг": ["Распил заготовки", "Токарка"],
            "Плита": ["Распил заготовки", "Фрезеровка"]
        }

        for option in ["Лист", "Круг", "Плита"]:
            rb = tk.Radiobutton(options_frame, text=option, value=option,
                                variable=self.radio_var, command=self.update_checkboxes)
            rb.pack(anchor=tk.W)

        # Чекбоксы операций
        self.checkbox_frame = tk.Frame(frame_right)
        self.checkbox_frame.pack(anchor=tk.W, padx=20, pady=10)

        # Кнопки
        btn_frame = tk.Frame(frame_right)
        btn_frame.pack(side="bottom", pady=20)

        self.btn_continue = tk.Button(btn_frame, text="Продолжить", state=tk.DISABLED,
                                      command=self.on_continue, font=("Arial", 12), width=15, bg="lightblue")
        self.btn_continue.pack(side="left", padx=10)

        self.btn_skip = tk.Button(btn_frame, text="Пропустить", command=self.on_skip,
                                  font=("Arial", 12), width=15, bg="lightgray")
        self.btn_skip.pack(side="left", padx=10)

        # Словарь описаний операций
        self.operation_descriptions = {
            "Вырубка заготовки": "Отделение заготовки по чертежу с припуском на дальнейшую обработку.",
            "Пробивка": "Пробивка отверстий в заготовке в соответствии с чертежом.",
            "Гибка": "Гибка заготовки в соответствии с чертежом.",
            "Распил заготовки": "Отделение заготовки по чертежу с припуском на дальнейшую обработку.",
            "Токарка": "Выполнение токарных операций в заготовке по чертежу.",
            "Сверление": "Выполнение отверстий в заготовке по чертежу.",
            "Нарезка резьбы": "Выполнение резьбы по чертежу.",
            "Фрезеровка": "Выполнение фрезеровальных операций в заготовке по чертежу."
        }

        # Инициализация чекбоксов
        self.check_boxes = []
        self.check_vars = {}
        self.update_checkboxes()

    def pdf_to_image(self, pdf_path, zoom=None):
        """Конвертирует первую страницу PDF в PIL-изображение и масштабирует"""
        zoom = zoom or self.zoom_level
        doc = fitz.open(pdf_path)
        page = doc.load_page(0)

        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        doc.close()

        return Image.open(io.BytesIO(pix.tobytes("png")))

    def update_preview(self):
        """Обновляет изображение после изменения масштаба"""
        try:
            self.img = self.pdf_to_image(self.pdf_path, self.zoom_level)
            self.tk_img_ref = ImageTk.PhotoImage(self.img)
            self.canvas.delete("all")
            self.image_on_canvas = self.canvas.create_image(0, 0, anchor="nw", image=self.tk_img_ref)
            self.canvas.config(scrollregion=self.canvas.bbox(tk.ALL))
        except Exception as e:
            print(f"[!] Не удалось обновить превью: {e}")

    def zoom_in(self):
        self.zoom_level *= 1.2
        self.update_preview()

    def zoom_out(self):
        self.zoom_level /= 1.2
        self.update_preview()

    def on_mouse_down(self, event):
        self.start_x = self.canvas.canvasx(event.x)
        self.start_y = self.canvas.canvasy(event.y)
        self.rect = self.canvas.create_rectangle(self.start_x, self.start_y, self.start_x, self.start_y, outline="red", width=2)

    def on_mouse_drag(self, event):
        cur_x = self.canvas.canvasx(event.x)
        cur_y = self.canvas.canvasy(event.y)
        self.canvas.coords(self.rect, self.start_x, self.start_y, cur_x, cur_y)

    def on_mouse_up(self, event):
        end_x = self.canvas.canvasx(event.x)
        end_y = self.canvas.canvasy(event.y)
        x1 = min(self.start_x, end_x)
        y1 = min(self.start_y, end_y)
        x2 = max(self.start_x, end_x)
        y2 = max(self.start_y, end_y)
        self.selected_image_area = (x1, y1, x2, y2)
        self.canvas.delete(self.rect)

    def update_checkboxes(self):
        selected_type = self.radio_var.get()
        ops = self.operations[selected_type]

        for cb in self.check_boxes:
            cb.destroy()
        self.check_boxes.clear()
        self.check_vars.clear()

        for op in ops:
            var = tk.BooleanVar(value=op in self.default_checks[selected_type])
            cb = tk.Checkbutton(self.checkbox_frame, text=op, variable=var, command=self.toggle_continue)
            cb.pack(anchor=tk.W)
            self.check_vars[op] = var
            self.check_boxes.append(cb)

        self.toggle_continue()

    def toggle_continue(self):
        if any(var.get() for var in self.check_vars.values()):
            self.btn_continue.config(state=tk.NORMAL)
        else:
            self.btn_continue.config(state=tk.DISABLED)

    def on_continue(self):
        self.выбранный_тип = self.radio_var.get()
        self.выбранные_операции = [op for op, var in self.check_vars.items() if var.get()]
        self.материал = self.material_entry.get().strip()

        if not self.выбранный_тип:
            messagebox.showwarning("Ошибка", "Тип заготовки не выбран!")
            return

        if not self.выбранные_операции:
            messagebox.showwarning("Ошибка", "Выберите хотя бы одну операцию!")
            return

        self.root.quit()
        self.root.destroy()

    def on_skip(self):
        self.выбранный_тип = None
        self.выбранные_операции = []
        self.root.quit()
        self.root.destroy()


def replace_all(document, old_text, new_text):
    replaced = False
    for para in document.paragraphs:
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, str(new_text or ""))
                replaced = True

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, str(new_text or ""))
                            replaced = True
    return replaced


def insert_image(document, keyword, image_path, width_inches=6):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Поиск в обычных параграфах
    for para in document.paragraphs:
        if keyword in para.text:
            print(f"[DEBUG] Найден {keyword} в параграфе")
            para.text = ""
            run = para.add_run()
            try:
                run.add_picture(image_path, width=Inches(width_inches))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                return True
            except Exception as e:
                print(f"[!] Не удалось вставить картинку в параграф: {e}")
                return False

    # Поиск в таблицах
    for table in document.tables:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if keyword in para.text:
                        print(f"[DEBUG] Найден {keyword} в таблице ({row_idx}, {cell_idx})")
                        para.text = ""
                        run = para.add_run()
                        try:
                            run.add_picture(image_path, width=Inches(width_inches))
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            return True
                        except Exception as e:
                            print(f"[!] Не удалось вставить картинку в таблицу: {e}")
                            return False

    print(f"[!] Метка '{keyword}' не найдена в документе")
    return False


def extract_and_save_image(pdf_path, output_path, area=None, zoom_preview=3, zoom_export=6):
    """
    Сохраняет область из PDF с масштабированием
    :param pdf_path: путь к PDF
    :param output_path: куда сохранить PNG
    :param area: координаты (x1, y1, x2, y2) на экране
    :param zoom_preview: масштаб превью
    :param zoom_export: масштаб экспорта
    """
    try:
        doc = fitz.open(pdf_path)
        page = doc.load_page(0)

        if area is None:
            clip_rect = page.rect
        else:
            orig_x1 = area[0] / zoom_preview
            orig_y1 = area[1] / zoom_preview
            orig_x2 = area[2] / zoom_preview
            orig_y2 = area[3] / zoom_preview
            clip_rect = fitz.Rect(orig_x1, orig_y1, orig_x2, orig_y2)

        mat_export = fitz.Matrix(zoom_export, zoom_export)
        pix_export = page.get_pixmap(matrix=mat_export, clip=clip_rect)
        pix_export.save(output_path, "png")
        doc.close()
        print(f"[OK] Изображение сохранено: {output_path}")
        return True
    except Exception as e:
        print(f"[!] Не удалось сохранить изображение: {e}")
        return False


# --- Основной цикл ---
for filename in os.listdir(folder_path):
    if filename.lower().endswith(".pdf"):
        file_path = os.path.join(folder_path, filename)

        try:
            doc_pdf = fitz.open(file_path)
            page = doc_pdf.load_page(0)

            # Определение формата листа
            is_large_sheet = page.rect.width > 800 or page.rect.height > 1000
            regions = rects_A3 if is_large_sheet else rects_A4

            # Извлечение данных из регионов
            ОбозначениеДетали = extract_text_from_rect(page, regions[0])
            НаименованиеДетали = extract_text_from_rect(page, regions[1])
            МатериалДетали = extract_text_from_rect(page, regions[2])
            РазработалДеталь = extract_text_from_rect(page, regions[3])
            ДатаЧертежа = extract_text_from_rect(page, regions[4])

            doc_pdf.saveIncr()
            doc_pdf.close()

            # Нормализация строк
            ОбозначениеДетали = normalize_dashes_and_spaces(ОбозначениеДетали)
            НаименованиеДетали = clean_filename(НаименованиеДетали)
            МатериалДетали = clean_filename(МатериалДетали)
            РазработалДеталь = clean_filename(РазработалДеталь)
            ДатаЧертежа = clean_filename(ДатаЧертежа)

            print(f"[DEBUG] {filename}:")
            print(f"Обозначение: {ОбозначениеДетали}")
            print(f"Наименование: {НаименованиеДетали}")
            print(f"Материал: {МатериалДетали}")
            print(f"Разработал: {РазработалДеталь}")
            print(f"Дата: {ДатаЧертежа}")

            # 🖼 Открываем GUI
            root = tk.Tk()
            app = App(root, file_path, ОбозначениеДетали, НаименованиеДетали, МатериалДетали, РазработалДеталь, ДатаЧертежа)

            try:
                root.mainloop()
            except Exception as e:
                print(f"[!] Ошибка GUI: {e}")
                continue

            if not app.выбранные_операции and not app.выбранный_тип:
                print(f"[SKIP] {filename} → пропущен.")
                continue

            # 📄 Создание .docx
            template_docx_path = os.path.join(folder_path, "template.docx")
            if not os.path.exists(template_docx_path):
                print(f"[!] Шаблон '{template_docx_path}' не найден → пропуск файла")
                continue

            base_name = os.path.splitext(filename)[0]
            docx_file = os.path.join(folder_path, f"{base_name}.docx")

            try:
                document = Document(template_docx_path)

                # Замена базовых полей
                replace_all(document, "#ОбозначениеДетали", app.обозначение)
                replace_all(document, "#НаименованиеДетали", app.наименование)
                replace_all(document, "#МатериалДетали", app.материал)
                replace_all(document, "#РазработалДеталь", app.разработал)
                replace_all(document, "#ДатаЧертежа", app.дата)
                replace_all(document, "#ТипЗаготовки", app.выбранный_тип or "")

                # Первая операция — получение материала
                replace_all(document, "#ОписаниеТехноОперации1",
                            "Получение материала со склада в соответствии с сменным заданием и чертежом.")

                # Последующие операции
                for idx, op in enumerate(app.выбранные_операции, start=2):
                    desc = app.operation_descriptions.get(op, f"Описание операции для {op} отсутствует.")
                    replace_all(document, f"#ОписаниеТехноОперации{idx}", desc)

                # Сохраняем изображение
                temp_image_path = os.path.join(folder_path, "temp_selected.png")
                area_to_use = app.selected_image_area if app.selected_image_area else (
                    0, 0, app.tk_img_ref.width(), app.tk_img_ref.height()
                )

                if extract_and_save_image(file_path, temp_image_path, area_to_use, zoom_preview=app.zoom_level, zoom_export=6):
                    if not insert_image(document, "#Картинка", temp_image_path):
                        print("[!] Не удалось вставить изображение в документ")
                else:
                    print("[!] Не удалось сохранить выделенную область")

                # ❌ Убрали очистку всех хеш-тегов, чтобы не повредить шаблон
                # remove_all_hash_tags(document)

                # Сохранение документа
                document.save(docx_file)
                print(f"[OK] Сохранён файл: {docx_file}")

            except Exception as e:
                print(f"[Ошибка при работе с Word]: {e}")
                continue

        except Exception as e:
            print(f"[Ошибка] Не удалось обработать '{filename}': {e}")
            try:
                doc_pdf.close()
            except:
                pass
            continue

messagebox.showinfo("Готово", "Обработка завершена.")
input("\nНажмите Enter для выхода...")